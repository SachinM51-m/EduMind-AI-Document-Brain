import os
import sys
from typing import List, Dict, Any
import PyPDF2
import docx

try:
    import win32com.client
    HAS_WIN32COM = True
except ImportError:
    HAS_WIN32COM = False

class RecursiveCharacterTextSplitter:
    """
    A custom implementation of a recursive character text splitter that splits text 
    by a list of separators recursively to maintain semantic chunk integrity.
    """
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200, separators: List[str] = None):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.separators = separators or ["\n\n", "\n", " ", ""]

    def split_text(self, text: str) -> List[str]:
        return self._split_text(text, self.separators)

    def _split_text(self, text: str, separators: List[str]) -> List[str]:
        if not text:
            return []
            
        if len(text) <= self.chunk_size:
            return [text]

        if not separators:
            # Fallback character-level split if we ran out of separators
            return [text[i:i + self.chunk_size] for i in range(0, len(text), self.chunk_size)]

        separator = separators[0]
        next_separators = separators[1:]

        if separator == "":
            splits = list(text)
        else:
            splits = text.split(separator)

        chunks = []
        current_chunk = []
        current_length = 0

        for split in splits:
            if len(split) > self.chunk_size:
                # If a single split is larger than the chunk size, split it recursively with remaining separators
                if current_chunk:
                    chunks.append(separator.join(current_chunk))
                    current_chunk = []
                    current_length = 0
                
                sub_splits = self._split_text(split, next_separators)
                chunks.extend(sub_splits)
            else:
                sep_len = len(separator) if current_chunk else 0
                if current_length + sep_len + len(split) > self.chunk_size:
                    # Save current chunk
                    if current_chunk:
                        chunks.append(separator.join(current_chunk))
                    
                    # Backtrack to build the overlap for the next chunk
                    overlap_chunk = []
                    overlap_len = 0
                    for item in reversed(current_chunk):
                        item_sep_len = len(separator) if overlap_chunk else 0
                        if overlap_len + item_sep_len + len(item) <= self.chunk_overlap:
                            overlap_chunk.insert(0, item)
                            overlap_len += item_sep_len + len(item)
                        else:
                            break
                    
                    current_chunk = overlap_chunk
                    current_length = overlap_len
                
                current_chunk.append(split)
                current_length += (len(separator) if len(current_chunk) > 1 else 0) + len(split)

        if current_chunk:
            chunks.append(separator.join(current_chunk))

        return [c.strip() for c in chunks if c.strip()]


def extract_text_from_pdf(file_path: str) -> str:
    """Extracts raw text from a PDF file using PyPDF2."""
    text = ""
    with open(file_path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        for page_idx, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text


def extract_text_from_docx(file_path: str) -> str:
    """Extracts raw text from a DOCX file using python-docx."""
    doc = docx.Document(file_path)
    text_parts = []
    
    # Extract text from paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)
            
    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    text_parts.append(cell_text)
                    
    return "\n".join(text_parts)


def extract_text_from_doc(file_path: str) -> str:
    """
    Extracts raw text from a legacy .doc file.
    Attempts to use MS Word COM interface on Windows, falls back to reporting manual conversion error.
    """
    if not HAS_WIN32COM:
        raise ImportError(
            "Extracting text from legacy .doc files on Windows requires the 'pywin32' library. "
            "Please convert the file to .docx or run: pip install pywin32"
        )
    
    import pythoncom
    pythoncom.CoInitialize()
    word = None
    doc = None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        abs_path = os.path.abspath(file_path)
        doc = word.Documents.Open(abs_path)
        text = doc.Content.Text
        return text
    except Exception as e:
        raise RuntimeError(f"Failed to read .doc file using MS Word COM interface: {str(e)}")
    finally:
        if doc:
            try:
                doc.Close(False)
            except Exception:
                pass
        if word:
            try:
                word.Quit()
            except Exception:
                pass


def process_document(file_path: str, chunk_size: int = 1000, chunk_overlap: int = 200) -> List[Dict[str, Any]]:
    """
    Processes a document: extracts its text, splits it into chunks, 
    and returns a list of dictionaries with content and metadata.
    """
    _, ext = os.path.splitext(file_path.lower())
    
    if ext == ".pdf":
        raw_text = extract_text_from_pdf(file_path)
    elif ext == ".docx":
        raw_text = extract_text_from_docx(file_path)
    elif ext == ".doc":
        raw_text = extract_text_from_doc(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}. Only PDF, DOCX, and DOC are supported.")
    
    if not raw_text.strip():
        return []
        
    splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    chunks = splitter.split_text(raw_text)
    
    processed_chunks = []
    for idx, chunk in enumerate(chunks):
        processed_chunks.append({
            "text": chunk,
            "metadata": {
                "source": os.path.basename(file_path),
                "chunk_index": idx,
                "total_chunks": len(chunks)
            }
        })
        
    return processed_chunks


if __name__ == "__main__":
    # Self-test/demo code
    if len(sys.argv) > 1:
        test_file = sys.argv[1]
        if os.path.exists(test_file):
            print(f"Processing test file: {test_file}")
            chunks = process_document(test_file)
            print(f"Generated {len(chunks)} chunks.")
            if chunks:
                print(f"First chunk snippet:\n{chunks[0]['text'][:200]}...")
        else:
            print(f"File not found: {test_file}")
    else:
        print("Usage: python document_processor.py <file_path>")
